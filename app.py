import streamlit as st
import pandas as pd
import io
import time
import json
import rag_engine as rag
import PyPDF2
import docx
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from fpdf import FPDF
import tempfile
import os

st.set_page_config(page_title="Agenzia AI Hub", layout="wide", page_icon="🚀")
st.markdown("""
    <style>
    .main-header {font-size: 2.2rem; font-weight: bold; color: #1E88E5; margin-bottom: 0.5rem;}
    .sub-header {font-size: 1.1rem; color: #555; margin-bottom: 1.5rem;}
    .debug-box {background-color: #282c34; color: #abb2bf; padding: 15px; border-radius: 5px; font-family: monospace; font-size: 0.85rem; white-space: pre-wrap;}
    .channel-box {background-color: #f0f2f6; padding: 10px; border-radius: 5px; margin-bottom: 5px;}
    .insight-box {background-color: #e8f4f8; border-left: 4px solid #1E88E5; padding: 15px; border-radius: 4px; margin-bottom: 15px;}
    .context-box {background-color: #fff3cd; border-left: 4px solid #ffc107; padding: 15px; border-radius: 4px; margin-bottom: 15px; font-size: 0.9rem;}
    </style>
""", unsafe_allow_html=True)

st.sidebar.title("🏢 Agenzia AI Hub")
all_clients = rag.get_all_clients()
client_options = ["➕ CREA NUOVO CLIENTE..."] + all_clients
selected_option = st.sidebar.selectbox("👤 Seleziona Cliente", client_options, key="client_selector")
client_id = ""

if selected_option == "➕ CREA NUOVO CLIENTE...":
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🆕 Nuovo Cliente")
    new_client_id = st.sidebar.text_input("ID Cliente (es. Nike, Mario_Rossi)", key="new_client_input").strip().replace(" ", "_")
    if st.sidebar.button("✅ CREA CLIENTE", type="primary", use_container_width=True):
        if not new_client_id:
            st.sidebar.error("⚠️ Inserisci un ID")
        elif new_client_id in all_clients:
            st.sidebar.warning(f"Il cliente '{new_client_id}' esiste già.")
        else:
            with st.sidebar.spinner(f"Creazione in corso..."):
                if rag.register_client(new_client_id):
                    rag.add_document(new_client_id, f"Cliente {new_client_id} inizializzato.", doc_type="sistema", source_file="sistema")
                    st.sidebar.success(f"✅ Cliente '{new_client_id}' creato!")
                    time.sleep(1)
                    st.rerun()
else:
    client_id = selected_option
    st.sidebar.markdown("---")
    st.sidebar.success(f"🟢 Cliente attivo: **{client_id}**")

if client_id:
    st.sidebar.markdown("---")
    with st.sidebar.expander("⚠️ Elimina Cliente"):
        st.warning(f"Stai per eliminare **TUTTI** i dati di '{client_id}'.")
        confirm_text = st.text_input(f"Per confermare, scrivi: {client_id}", key="delete_confirm")
        if st.button(f"🗑️ ELIMINA '{client_id}'", type="secondary", use_container_width=True):
            if confirm_text.strip() == client_id:
                with st.spinner("Eliminazione in corso..."):
                    success, message = rag.delete_client(client_id)
                    if success:
                        st.success(f"✅ Eliminato.")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error(f"❌ Errore: {message}")
            else:
                st.error("❌ Testo non corrispondente.")

if not client_id:
    st.markdown('<div class="main-header">Benvenuto in Agenzia AI Hub</div>', unsafe_allow_html=True)
    st.info("👈 Seleziona o crea un cliente per iniziare.")
    st.stop()

st.markdown(f'<div class="main-header">Dashboard: {client_id}</div>', unsafe_allow_html=True)
task_type = st.radio("🤖 Scegli l'Agente", [
    "🧠 Carica e Gestisci Memoria", 
    "📅 Piano Editoriale Completo", 
    "📊 Report ADS Performance",
    "📱 Report Social Organico",
    "🔍 Analisi Competitor / Trend"
], horizontal=True)
st.markdown("---")

if task_type == "🧠 Carica e Gestisci Memoria":
    st.markdown('<div class="sub-header">Alimenta o modifica la memoria strategica del cliente</div>', unsafe_allow_html=True)
    st.markdown("### 📂 Memoria Attuale")
    memory_summary = rag.get_memory_summary(client_id)
    reverse_mapping = {
        "brand_book": "📘 Brand Book / Linee Guida", "icp_personas": "👤 ICP / Personas & Pain/Gain",
        "gestione_obiezioni": "🛡️ Gestione Obiezioni", "esempi_copy": "✍️ Esempi di Copy Approvati",
        "istruzioni_creazione": "📝 Istruzioni Specifiche di Creazione", "note_call": "📞 Note da Call / Briefing",
        "regole_negative": "🚫 Regole Negative", "report_dati": "📊 Report / Dati Precedenti",
        "link_riferimento": "🔗 Link Asset, Competitor e Fonti", "sistema": "⚙️ Sistema",
        "regola_stile": "🧠 Regole Apprese (Opzione B)", "errore": "❌ Errore DB"
    }
    has_data = any(key != "errore" for key in memory_summary)
    if not has_data:
        st.info("La memoria di questo cliente è vuota.")
    else:
        st.success("✅ Memoria caricata correttamente.")
        for doc_type, data in memory_summary.items():
            if doc_type == "errore": continue
            display_name = reverse_mapping.get(doc_type, f"📁 {doc_type}")
            with st.expander(f"**{display_name}** ({data.get('count', 0)} blocchi)"):
                for f in data.get("files", []):
                    col_f1, col_f2 = st.columns([4, 1])
                    with col_f1: st.markdown(f"- 📄 `{f}`")
                    with col_f2:
                        if st.button("🗑️", key=f"del_{doc_type}_{f}"):
                            success, msg = rag.delete_specific_file(client_id, doc_type, f)
                            if success: st.success(msg); time.sleep(1); st.rerun()
                if st.button(f"🗑️ ELIMINA TUTTA LA CATEGORIA", key=f"del_cat_{doc_type}", type="secondary"):
                    success, msg = rag.delete_category(client_id, doc_type)
                    if success: st.success(msg); time.sleep(1.5); st.rerun()
    
    st.markdown("---")
    st.markdown("### ➕ Aggiungi Nuovo Contenuto")
    col1, col2 = st.columns([2, 1])
    with col1:
        uploaded_files = st.file_uploader("📎 Carica file (PDF, DOCX, TXT, CSV)", type=["pdf", "docx", "txt", "csv"], accept_multiple_files=True)
        manual_text = st.text_area("Oppure incolla testo manuale:", height=150)
        url_input = st.text_input("🌐 Incolla URL da scansionare (uno alla volta)", placeholder="https://...")
        if st.button("🌐 Scansiona e Salva Link", type="secondary"):
            if url_input.strip():
                with st.spinner(f"Scansione di {url_input} in corso..."):
                    success, msg = rag.scrape_and_save_url(client_id, url_input.strip(), doc_type="link_riferimento")
                    if success: st.success(msg); time.sleep(1.5); st.rerun()
                    else: st.error(msg)
            else: st.warning("Inserisci un URL valido.")
    with col2:
        standard_categories = ["📘 Brand Book / Linee Guida", "👤 ICP / Personas & Pain/Gain", "🛡️ Gestione Obiezioni", "✍️ Esempi di Copy Approvati", "📝 Istruzioni Specifiche di Creazione", "📞 Note da Call / Briefing", "🚫 Regole Negative", "📊 Report / Dati Precedenti", "🔗 Link Asset, Competitor e Fonti", "➕ Scrivi una categoria personalizzata..."]
        selected_cat = st.selectbox("Scegli", standard_categories, key="cat_select")
        if "➕" in selected_cat:
            doc_type = st.text_input("Nome categoria", key="custom_cat").strip().lower().replace(" ", "_")
        else:
            mapping = {"📘 Brand Book / Linee Guida": "brand_book", "👤 ICP / Personas & Pain/Gain": "icp_personas", "🛡️ Gestione Obiezioni": "gestione_obiezioni", "✍️ Esempi di Copy Approvati": "esempi_copy", "📝 Istruzioni Specifiche di Creazione": "istruzioni_creazione", "📞 Note da Call / Briefing": "note_call", "🚫 Regole Negative": "regole_negative", "📊 Report / Dati Precedenti": "report_dati", "🔗 Link Asset, Competitor e Fonti": "link_riferimento"}
            doc_type = mapping.get(selected_cat, "generico")
        st.info(f"Salverai come: **`{doc_type}`**")

    if st.button("💾 Salva nella Memoria", type="primary"):
        debug_info = []
        if uploaded_files:
            for uploaded_file in uploaded_files:
                try:
                    extracted = ""
                    uploaded_file.seek(0)
                    if uploaded_file.type in ["text/plain", "text/csv"]: extracted = uploaded_file.read().decode("utf-8")
                    elif uploaded_file.type == "application/pdf":
                        reader = PyPDF2.PdfReader(uploaded_file)
                        extracted = "\n".join([page.extract_text() or "" for page in reader.pages])
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = docx.Document(uploaded_file)
                        extracted = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    if len(extracted.strip()) > 0:
                        ok, msg = rag.add_document(client_id, extracted, doc_type, source_file=uploaded_file.name)
                        debug_info.append(msg)
                except Exception as e: debug_info.append(f"❌ {uploaded_file.name}: {str(e)}")
        if manual_text.strip():
            ok, msg = rag.add_document(client_id, manual_text.strip(), doc_type, source_file="testo_manuale")
            debug_info.append(msg)
        if debug_info:
            st.markdown(f'<div class="debug-box">' + "\n".join(debug_info) + "</div>", unsafe_allow_html=True)
            time.sleep(1.5); st.rerun()

elif task_type == "📅 Piano Editoriale Completo":
    st.markdown('<div class="sub-header">Generazione professionale: 3 contenuti alla volta, copy profondi e strutturati</div>', unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1: mese = st.text_input("📅 Periodo", "Gennaio 2025")
    with col2: obiettivo = st.selectbox("🎯 Obiettivo", ["Brand Awareness", "Lead Generation", "Lancio Prodotto", "Fidelizzazione", "Engagement"])
    with col3: durata = st.selectbox("⏱️ Durata", ["1 settimana", "2 settimane", "1 mese", "3 mesi"])
    tema = st.text_input("💡 Tema Centrale", "Es. Posizionamento come esperti")
    istruzioni_extra = st.text_area("📝 Note specifiche", placeholder="Es. Evitare la parola X, enfatizzare Y...")
    
    st.markdown("---")
    st.markdown("### 📱 2. Contenuti Social")
    canali_config = {}
    canali_disponibili = {"LinkedIn": "Post LinkedIn", "Instagram_Feed": "Post IG Feed", "Instagram_Reels": "Reel IG", "Instagram_Stories": "Story IG", "Facebook": "Post FB", "TikTok": "TikTok"}
    cols = st.columns(2)
    for idx, (canale, label) in enumerate(canali_disponibili.items()):
        with cols[idx % 2]:
            quant = st.slider(f"**{label}**", 0, 20, 0, key=f"slider_{canale}")
            if quant > 0: canali_config[canale] = quant

    totale = sum(canali_config.values())
    if totale == 0:
        st.warning("⚠️ Configura almeno un contenuto per generare il piano.")
    else:
        st.success(f"🎯 Piano configurato: {totale} contenuti totali (verranno generati in batch di 3 per garantire qualità)")
        if st.button(f"🚀 GENERA PRIMI 3 CONTENUTI", type="primary", use_container_width=True):
            with st.spinner("Recupero contesto e scrittura copy professionali..."):
                context = rag.get_client_context(client_id, "brand book, ICP, personas, pain, gain, obiezioni, istruzioni di creazione, tono di voce, link riferimento")
                
                with st.expander("🔍 Vedi cosa sta leggendo l'AI (Contesto Recuperato dai tuoi file)", expanded=False):
                    st.markdown(f'<div class="context-box">{context}</div>', unsafe_allow_html=True)
                
                canali_list = list(canali_config.keys())[:3] # Prende i primi 3 configurati per il batch
                
                prompt_pe = (
                    f"Sei un Senior Copywriter e Content Strategist. Genera ESATTAMENTE 3 contenuti in formato JSON STRICT.\n\n"
                    f"## CONTESTO CLIENTE (BASE ASSOLUTA):\n{context}\n\n"
                    f"## CONFIGURAZIONE:\nPeriodo: {mese} | Obiettivo: {obiettivo} | Tema Centrale: {tema} | Note: {istruzioni_extra}\n"
                    f"Canali da usare in questo batch: {', '.join(canali_list)}\n\n"
                    f"## FORMATO OUTPUT JSON:\n[{{\"tipo\": \"...\", \"data\": \"YYYY-MM-DD\", \"titolo\": \"...\", \"hook\": \"...\", \"copy\": \"...\", \"cta\": \"...\", \"brief_visivo\": \"...\", \"hashtag_seo\": \"...\", \"note\": \"...\"}}, ...]\n\n"
                    f"## VINCOLI OBBLIGATORI:\n"
                    f"1. **LUNGHEZZA COPY**: LinkedIn = 150-250 parole. Instagram = 80-150 parole. TikTok/Reel = Script di 150-300 caratteri. Vietato testi brevi di 1-2 righe.\n"
                    f"2. **STRUTTURA COPY**: Intro (Hook) → 2-3 paragrafi di sviluppo (pain, soluzione, prova sociale) → CTA chiara. Usa a capo reali.\n"
                    f"3. **ALLINEAMENTO TEMA**: Ogni contenuto deve ruotare ESPlicitamente attorno a: '{tema}'.\n"
                    f"4. **VIETATO**: Placeholder, 'Lorem ipsum', frasi generiche, ripetizioni. Scrivi testi PRONTI ALLA PUBBLICAZIONE.\n"
                    f"5. Rispondi SOLO con il JSON valido, senza markdown o testo extra."
                )
                response = rag.llm.invoke(prompt_pe).content
                
                try:
                    # Pulizia JSON
                    clean_json = response.replace("```json", "").replace("```", "").strip()
                    data_list = json.loads(clean_json)
                    df = pd.DataFrame(data_list)
                    st.session_state['ped_batch'] = df
                    st.success("✅ 3 contenuti generati con successo! Copy profondi e strutturati.")
                    st.data_editor(df, num_rows="dynamic", key="editable_df", height=600, use_container_width=True)
                    
                    # DOCX Export
                    doc = docx.Document()
                    doc.add_heading(f'Piano Editoriale: {client_id}', 0)
                    doc.add_paragraph(f'Periodo: {mese} | Obiettivo: {obiettivo} | Tema: {tema}')
                    doc.add_paragraph('_' * 50)
                    for idx, row in df.iterrows():
                        doc.add_heading(f"{idx + 1}. {row.get('titolo', 'Contenuto')}", level=1)
                        doc.add_paragraph(f"📌 Tipo: {row.get('tipo')} | 📅 Data: {row.get('data')}", style='Normal')
                        if row.get('hook'): doc.add_paragraph(f"🎣 Hook: {row['hook']}")
                        if row.get('copy'): 
                            doc.add_paragraph("📝 Copy:", style='Heading 3')
                            for p in str(row['copy']).split('\n'): doc.add_paragraph(p.strip())
                        if row.get('cta'): doc.add_paragraph(f"👉 CTA: {row['cta']}")
                        if row.get('brief_visivo'): doc.add_paragraph(f"🎨 Brief: {row['brief_visivo']}")
                        doc.add_paragraph("__________________________________________________________________________________________")
                    
                    buf = io.BytesIO()
                    doc.save(buf); buf.seek(0)
                    st.download_button("📥 Scarica WORD", buf, f"PED_{client_id}_{mese.replace(' ','_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    
                except Exception as e:
                    st.error("⚠️ Errore parsing JSON. Output grezzo:"); st.code(response)

    if 'ped_batch' in st.session_state:
        st.markdown("---")
        if st.button("💾 Salva e Insegna (Opzione B)", type="secondary"):
            with st.spinner("Analisi correzioni in corso..."):
                result = rag.save_and_teach(client_id, json.dumps(st.session_state['ped_batch'].to_dict(orient='records')), st.session_state['editable_df'].to_json(orient='records'))
                st.success(result)

elif task_type == "📊 Report ADS Performance":
    st.markdown('<div class="sub-header">Genera report PDF professionale con analisi performance ADS</div>', unsafe_allow_html=True)
    st.info("💡 **Istruzioni:** Scarica il report CSV da Meta Ads Manager o Google Ads.")
    col1, col2 = st.columns([1, 2])
    with col1:
        st.markdown("### 1. Dati Campagna")
        uploaded_report = st.file_uploader("📎 Carica Report CSV", type=["csv"], key="ads_upload")
        date_range = st.text_input("📅 Intervallo Date", "Es. 1-31 Ottobre 2024")
        obiettivo_campagna = st.selectbox("🎯 Obiettivo", ["Lead Generation", "Vendite/E-commerce", "Brand Awareness", "Traffico Sito"])
    with col2:
        st.markdown("### 2. Anteprima Contesto")
        context_preview = rag.get_client_context(client_id, "ICP, obiettivi strategici, pain gain")
        st.markdown(f'<div class="debug-box" style="max-height: 200px; overflow-y: auto; font-size: 0.75rem;">{context_preview[:500]}...</div>', unsafe_allow_html=True)

    if uploaded_report is not None:
        if st.button("🚀 Genera Report PDF", type="primary", use_container_width=True):
            with st.spinner("Analisi dati e generazione report in corso..."):
                try:
                    df_report = pd.read_csv(uploaded_report)
                    df_report.columns = [col.strip().lower().replace(' ', '_') for col in df_report.columns]
                    spend_col = next((c for c in df_report.columns if 'spesa' in c or 'spend' in c or 'cost' in c), None)
                    impr_col = next((c for c in df_report.columns if 'impression' in c), None)
                    click_col = next((c for c in df_report.columns if 'click' in c and 'ctr' not in c), None)
                    ctr_col = next((c for c in df_report.columns if 'ctr' in c), None)
                    cpa_col = next((c for c in df_report.columns if 'cpa' in c or 'cost_per' in c), None)
                    roas_col = next((c for c in df_report.columns if 'roas' in c or 'return' in c), None)
                    conv_col = next((c for c in df_report.columns if 'conv' in c or 'conversion' in c), None)
                    campaign_col = next((c for c in df_report.columns if 'camp' in c or 'name' in c), None)
                    
                    total_spend = df_report[spend_col].sum() if spend_col else 0
                    total_impr = df_report[impr_col].sum() if impr_col else 0
                    total_click = df_report[click_col].sum() if click_col else 0
                    total_conv = df_report[conv_col].sum() if conv_col else 0
                    avg_ctr = (total_click / total_impr * 100) if total_impr > 0 else 0
                    avg_cpa = (total_spend / total_conv) if total_conv > 0 else 0
                    avg_roas = df_report[roas_col].mean() if roas_col else 0
                    
                    csv_sample = df_report.head(20).to_string()
                    prompt_analysis = (
                        f"Sei un Senior Media Buyer. Analizza questo report ADS e genera un commento strategico in italiano.\n\n"
                        f"## DATI:\n{csv_sample}\n\n"
                        f"## CONTESTO:\nCliente: {client_id}\nObiettivo: {obiettivo_campagna}\nPeriodo: {date_range}\n"
                        f"Metriche aggregate: Spesa={total_spend:.2f}, Impressioni={total_impr}, Click={total_click}, CTR={avg_ctr:.2f}%, CPA={avg_cpa:.2f}, ROAS={avg_roas:.2f}\n\n"
                        f"## COMPITO:\nRispondi in 4 sezioni concise (max 150 parole ciascuna):\n"
                        f"1. **Executive Summary**: Performance complessiva e raggiungimento obiettivo\n"
                        f"2. **Cosa ha Funzionato**: Pattern vincenti (formati, copy, audience)\n"
                        f"3. **Anomalie e Sprechi**: Dove si brucia budget\n"
                        f"4. **Raccomandazioni**: 3 azioni concrete per il prossimo periodo"
                    )
                    ai_analysis = rag.llm.invoke(prompt_analysis).content
                    
                    pdf = FPDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 24)
                    pdf.cell(0, 60, "", ln=True)
                    pdf.cell(0, 20, f"Report Performance ADS", ln=True, align="C")
                    pdf.set_font("Arial", "", 16)
                    pdf.cell(0, 15, f"Cliente: {client_id}", ln=True, align="C")
                    pdf.cell(0, 10, f"Periodo: {date_range}", ln=True, align="C")
                    pdf.cell(0, 10, f"Obiettivo: {obiettivo_campagna}", ln=True, align="C")
                    pdf.set_font("Arial", "I", 10)
                    pdf.cell(0, 40, "", ln=True)
                    pdf.cell(0, 10, f"Generato il: {time.strftime('%d/%m/%Y')}", ln=True, align="C")
                    
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 18)
                    pdf.cell(0, 15, "Metriche Chiave", ln=True)
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 12)
                    pdf.set_fill_color(230, 230, 230)
                    pdf.cell(95, 10, "Metrica", border=1, fill=True)
                    pdf.cell(95, 10, "Valore", border=1, fill=True, ln=True)
                    pdf.set_font("Arial", "", 11)
                    metrics = [("Spesa Totale", f"€ {total_spend:.2f}"), ("Impressioni", f"{total_impr:,.0f}"), ("Click Totali", f"{total_click:,.0f}"), ("CTR Medio", f"{avg_ctr:.2f}%"), ("Conversioni", f"{total_conv:,.0f}"), ("CPA Medio", f"€ {avg_cpa:.2f}"), ("ROAS Medio", f"{avg_roas:.2f}x")]
                    for metric, value in metrics:
                        pdf.cell(95, 8, metric, border=1)
                        pdf.cell(95, 8, value, border=1, ln=True)
                    
                    pdf.ln(10)
                    pdf.set_font("Arial", "B", 14)
                    pdf.cell(0, 10, "Visualizzazione Dati", ln=True)
                    
                    if campaign_col and spend_col:
                        fig, ax = plt.subplots(figsize=(8, 4))
                        top_campaigns = df_report.groupby(campaign_col)[spend_col].sum().nlargest(5)
                        top_campaigns.plot(kind='bar', ax=ax, color='#1E88E5')
                        ax.set_title('Top 5 Campagne per Spesa', fontsize=12, fontweight='bold')
                        ax.set_ylabel('Spesa (€)')
                        plt.xticks(rotation=45, ha='right')
                        plt.tight_layout()
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            plt.savefig(tmp.name, dpi=150, bbox_inches='tight')
                            tmp_path = tmp.name
                        plt.close()
                        pdf.image(tmp_path, x=10, w=190)
                        os.unlink(tmp_path)
                    
                    if len(df_report) > 1 and ctr_col and cpa_col:
                        pdf.add_page()
                        fig, axes = plt.subplots(1, 2, figsize=(10, 4))
                        df_report[ctr_col].hist(ax=axes[0], bins=20, color='#4CAF50', alpha=0.7)
                        axes[0].set_title('Distribuzione CTR', fontweight='bold')
                        df_report[cpa_col].hist(ax=axes[1], bins=20, color='#FF9800', alpha=0.7)
                        axes[1].set_title('Distribuzione CPA', fontweight='bold')
                        plt.tight_layout()
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            plt.savefig(tmp.name, dpi=150, bbox_inches='tight')
                            tmp_path = tmp.name
                        plt.close()
                        pdf.image(tmp_path, x=10, w=190)
                        os.unlink(tmp_path)
                    
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 18)
                    pdf.cell(0, 15, "Analisi Strategica AI", ln=True)
                    pdf.ln(5)
                    pdf.set_font("Arial", "", 10)
                    for line in ai_analysis.split('\n'):
                        clean_line = line.replace('**', '').replace('*', '')
                        if clean_line.strip():
                            pdf.multi_cell(0, 6, clean_line)
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                        pdf.output(tmp.name)
                        tmp_path = tmp.name
                    with open(tmp_path, 'rb') as pdf_file:
                        pdf_bytes = pdf_file.read()
                    
                    st.success("✅ Report generato con successo!")
                    st.markdown("### 📈 Anteprima Analisi")
                    st.markdown(f'<div class="insight-box">{ai_analysis}</div>', unsafe_allow_html=True)
                    st.download_button(label="📥 Scarica Report PDF", data=pdf_bytes, file_name=f"Report_ADS_{client_id}_{date_range.replace(' ', '_')}.pdf", mime="application/pdf")
                    os.unlink(tmp_path)
                except Exception as e:
                    st.error(f"❌ Errore: {str(e)}")

elif task_type == "📱 Report Social Organico":
    st.markdown('<div class="sub-header">Genera report PDF con analisi performance organica social</div>', unsafe_allow_html=True)
    st.info("💡 **Istruzioni:** Carica un CSV con i dati dei post social (esportato da Meta Business Suite, LinkedIn Analytics, ecc.).")
    col1, col2 = st.columns([1, 2])
    with col1:
        st.markdown("### 1. Dati Social")
        uploaded_social = st.file_uploader("📎 Carica CSV Post Social", type=["csv"], key="social_upload")
        date_range_social = st.text_input("📅 Intervallo Date", "Es. Ottobre 2024")
        piattaforme = st.multiselect("📱 Piattaforme Incluse", ["Instagram", "Facebook", "LinkedIn", "TikTok", "Twitter"], default=["Instagram", "LinkedIn"])
    with col2:
        st.markdown("### 2. Contesto Strategico")
        context_preview = rag.get_client_context(client_id, "ICP, tono di voce, obiettivi social")
        st.markdown(f'<div class="debug-box" style="max-height: 200px; overflow-y: auto; font-size: 0.75rem;">{context_preview[:500]}...</div>', unsafe_allow_html=True)

    if uploaded_social is not None:
        if st.button("🚀 Genera Report Social PDF", type="primary", use_container_width=True):
            with st.spinner("Analisi contenuti e generazione report..."):
                try:
                    df_social = pd.read_csv(uploaded_social)
                    df_social.columns = [col.strip().lower().replace(' ', '_') for col in df_social.columns]
                    platform_col = next((c for c in df_social.columns if 'piatt' in c or 'platform' in c), None)
                    content_type_col = next((c for c in df_social.columns if 'tipo' in c or 'type' in c), None)
                    likes_col = next((c for c in df_social.columns if 'like' in c or 'reaction' in c), None)
                    comments_col = next((c for c in df_social.columns if 'comment' in c), None)
                    shares_col = next((c for c in df_social.columns if 'share' in c or 'condiv' in c), None)
                    reach_col = next((c for c in df_social.columns if 'reach' in c or 'portata' in c), None)
                    engagement_col = next((c for c in df_social.columns if 'engagement' in c or 'eng_rate' in c), None)
                    
                    total_posts = len(df_social)
                    total_likes = df_social[likes_col].sum() if likes_col else 0
                    total_comments = df_social[comments_col].sum() if comments_col else 0
                    total_shares = df_social[shares_col].sum() if shares_col else 0
                    total_reach = df_social[reach_col].sum() if reach_col else 0
                    avg_engagement = df_social[engagement_col].mean() if engagement_col else 0
                    
                    if engagement_col: top_posts = df_social.nlargest(5, engagement_col)
                    elif likes_col: top_posts = df_social.nlargest(5, likes_col)
                    else: top_posts = df_social.head(5)
                    
                    social_sample = df_social.head(20).to_string()
                    prompt_social = (
                        f"Sei un Social Media Manager esperto. Analizza questo report di contenuti organici e genera un commento strategico in italiano.\n\n"
                        f"## DATI:\n{social_sample}\n\n"
                        f"## CONTESTO:\nCliente: {client_id}\nPeriodo: {date_range_social}\nPiattaforme: {', '.join(piattaforme)}\n"
                        f"Metriche: Post={total_posts}, Like={total_likes}, Commenti={total_comments}, Condivisioni={total_shares}, Reach={total_reach}, Engagement Rate Medio={avg_engagement:.2f}%\n\n"
                        f"## COMPITO:\nRispondi in 4 sezioni concise (max 150 parole ciascuna):\n"
                        f"1. **Performance Generale**: Crescita, engagement, reach e confronto con obiettivi\n"
                        f"2. **Contenuti Vincenti**: Quali formati, temi, copy hanno performato meglio e perché\n"
                        f"3. **Aree di Miglioramento**: Cosa non ha funzionato e opportunità mancata\n"
                        f"4. **Strategia Prossimo Periodo**: 3-4 raccomandazioni concrete per migliorare"
                    )
                    ai_analysis_social = rag.llm.invoke(prompt_social).content
                    
                    pdf = FPDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 24)
                    pdf.cell(0, 60, "", ln=True)
                    pdf.cell(0, 20, f"Report Performance Social Organico", ln=True, align="C")
                    pdf.set_font("Arial", "", 16)
                    pdf.cell(0, 15, f"Cliente: {client_id}", ln=True, align="C")
                    pdf.cell(0, 10, f"Periodo: {date_range_social}", ln=True, align="C")
                    pdf.cell(0, 10, f"Piattaforme: {', '.join(piattaforme)}", ln=True, align="C")
                    pdf.set_font("Arial", "I", 10)
                    pdf.cell(0, 40, "", ln=True)
                    pdf.cell(0, 10, f"Generato il: {time.strftime('%d/%m/%Y')}", ln=True, align="C")
                    
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 18)
                    pdf.cell(0, 15, "Metriche Chiave", ln=True)
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 12)
                    pdf.set_fill_color(230, 230, 230)
                    pdf.cell(95, 10, "Metrica", border=1, fill=True)
                    pdf.cell(95, 10, "Valore", border=1, fill=True, ln=True)
                    pdf.set_font("Arial", "", 11)
                    metrics_social = [("Post Totali", f"{total_posts}"), ("Like Totali", f"{total_likes:,.0f}"), ("Commenti", f"{total_comments:,.0f}"), ("Condivisioni", f"{total_shares:,.0f}"), ("Reach Totale", f"{total_reach:,.0f}"), ("Engagement Rate Medio", f"{avg_engagement:.2f}%")]
                    for metric, value in metrics_social:
                        pdf.cell(95, 8, metric, border=1)
                        pdf.cell(95, 8, value, border=1, ln=True)
                    
                    pdf.ln(10)
                    pdf.set_font("Arial", "B", 14)
                    pdf.cell(0, 10, "Visualizzazione Dati", ln=True)
                    
                    if platform_col and likes_col:
                        fig, ax = plt.subplots(figsize=(8, 4))
                        platform_perf = df_social.groupby(platform_col)[likes_col].sum()
                        platform_perf.plot(kind='bar', ax=ax, color=['#E1306C', '#1877F2', '#0A66C2', '#000000', '#1DA1F2'][:len(platform_perf)])
                        ax.set_title('Performance per Piattaforma (Like)', fontsize=12, fontweight='bold')
                        ax.set_ylabel('Like Totali')
                        plt.xticks(rotation=45, ha='right')
                        plt.tight_layout()
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            plt.savefig(tmp.name, dpi=150, bbox_inches='tight')
                            tmp_path = tmp.name
                        plt.close()
                        pdf.image(tmp_path, x=10, w=190)
                        os.unlink(tmp_path)
                    
                    if content_type_col and engagement_col:
                        pdf.add_page()
                        fig, ax = plt.subplots(figsize=(8, 4))
                        type_perf = df_social.groupby(content_type_col)[engagement_col].mean().nlargest(5)
                        type_perf.plot(kind='barh', ax=ax, color='#4CAF50')
                        ax.set_title('Top 5 Tipi Contenuto per Engagement', fontsize=12, fontweight='bold')
                        ax.set_xlabel('Engagement Rate (%)')
                        plt.tight_layout()
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                            plt.savefig(tmp.name, dpi=150, bbox_inches='tight')
                            tmp_path = tmp.name
                        plt.close()
                        pdf.image(tmp_path, x=10, w=190)
                        os.unlink(tmp_path)
                    
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 18)
                    pdf.cell(0, 15, "Top 5 Post del Periodo", ln=True)
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 9)
                    pdf.set_fill_color(230, 230, 230)
                    pdf.cell(60, 8, "Data/Tipo", border=1, fill=True)
                    pdf.cell(80, 8, "Contenuto (anteprima)", border=1, fill=True)
                    pdf.cell(25, 8, "Like", border=1, fill=True)
                    pdf.cell(25, 8, "Eng.", border=1, fill=True, ln=True)
                    pdf.set_font("Arial", "", 8)
                    for _, post in top_posts.iterrows():
                        data_tipo = str(post.get('data', post.get('date', 'N/A')))[:10]
                        if content_type_col: data_tipo += f" | {str(post.get(content_type_col, ''))[:15]}"
                        content_preview = str(post.get('testo', post.get('copy', post.get('content', 'N/A'))))[:40]
                        likes = f"{post.get(likes_col, 0):,.0f}" if likes_col else "N/A"
                        eng = f"{post.get(engagement_col, 0):.2f}%" if engagement_col else "N/A"
                        pdf.cell(60, 6, data_tipo, border=1)
                        pdf.cell(80, 6, content_preview, border=1)
                        pdf.cell(25, 6, likes, border=1)
                        pdf.cell(25, 6, eng, border=1, ln=True)
                    
                    pdf.add_page()
                    pdf.set_font("Arial", "B", 18)
                    pdf.cell(0, 15, "Analisi Strategica AI", ln=True)
                    pdf.ln(5)
                    pdf.set_font("Arial", "", 10)
                    for line in ai_analysis_social.split('\n'):
                        clean_line = line.replace('**', '').replace('*', '')
                        if clean_line.strip(): pdf.multi_cell(0, 6, clean_line)
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                        pdf.output(tmp.name)
                        tmp_path = tmp.name
                    with open(tmp_path, 'rb') as pdf_file:
                        pdf_bytes = pdf_file.read()
                    
                    st.success("✅ Report Social generato con successo!")
                    st.markdown("### 📈 Anteprima Analisi")
                    st.markdown(f'<div class="insight-box">{ai_analysis_social}</div>', unsafe_allow_html=True)
                    st.download_button(label="📥 Scarica Report Social PDF", data=pdf_bytes, file_name=f"Report_Social_{client_id}_{date_range_social.replace(' ', '_')}.pdf", mime="application/pdf")
                    os.unlink(tmp_path)
                except Exception as e:
                    st.error(f"❌ Errore: {str(e)}")

elif task_type == "🔍 Analisi Competitor / Trend":
    st.markdown('<div class="sub-header">Ricerca sul web trend o competitor</div>', unsafe_allow_html=True)
    query_ricerca = st.text_input("🔍 Cosa cercare? (es. 'trend marketing B2B gennaio 2025')")
    if st.button("🌐 Avvia Ricerca Web", type="primary"):
        if query_ricerca:
            with st.spinner("Ricerca in corso..."):
                search_results = rag.web_search(query_ricerca, num_results=5)
                if "Errore" in search_results or "⚠️" in search_results:
                    st.warning(search_results)
                else:
                    st.markdown("### 📊 Risultati")
                    st.markdown(search_results)
                    with st.spinner("Sintesi insight..."):
                        context = rag.get_client_context(client_id, "obiettivi strategici, ICP, pain gain, link competitor")
                        prompt_sintesi = f"Sei un Brand Strategist. Ricerca:\n{search_results}\nCliente: '{client_id}'. Contesto: {context}\nSintetizza in 3 insight pratici per il prossimo piano editoriale."
                        st.info(rag.llm.invoke(prompt_sintesi).content)